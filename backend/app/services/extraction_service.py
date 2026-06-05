import pandas as pd
import fitz  # PyMuPDF
import pytesseract
import re
import os
import time
from PIL import Image, ImageFilter, ImageEnhance
import numpy as np
from docx import Document
import zipfile
import io
from PIL import Image as PILImage
import openpyxl

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"



# EXCEL EXTRACTION — Week 6 + multi-sheet + your two fields
"""
def extract_excel(path: str):
    start = time.time()
    xls = pd.ExcelFile(path)
    sheets_output = []
    rows_per_sheet = {}
    columns_per_sheet = {}

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet_name, header=None)
        df = df.fillna("").astype(str)

        rows_per_sheet[sheet_name] = df.shape[0]
        columns_per_sheet[sheet_name] = df.shape[1]

        sheets_output.append({
            "sheet_name": sheet_name,
            "table_index": len(sheets_output),
            "headers": [],
            "rows": df.values.tolist()
        })

    metadata = {
        "sheet_count": len(xls.sheet_names),
        "rows_per_sheet": rows_per_sheet,
        "columns_per_sheet": columns_per_sheet,
        "file_size_kb": os.path.getsize(path) / 1024,
        "extraction_time_ms": int((time.time() - start) * 1000)
    }

    return {
        "tables": sheets_output,
        "metadata": metadata
    }

"""


def extract_excel(file_path: str):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    tables = []
    metadata = {"sheets_extracted": 0}

    
    def detect_header_row(rows):
        for i, row in enumerate(rows):
            # Must have at least 3 non-empty cells
            non_empty = [str(c).strip() for c in row if str(c).strip()]
            if len(non_empty) < 3:
                continue

            # Header row must NOT contain currency values
            #if any(re.match(r"^[\d\.,]+$", c) for c in non_empty):
            #    continue

            if any("$" in c or "€" in c or "£" in c for c in non_empty):
                continue

            # Next row must contain at least one numeric value
            if i + 1 < len(rows):
                next_row = rows[i + 1]
                next_non_empty = [str(c).strip() for c in next_row if str(c).strip()]

                # Accept mixed text + numeric
                if any(re.match(r"^[\d\.,]+$", c) for c in next_non_empty):
                    return i

        return None




    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Read all rows
        raw_rows = []
        for row in ws.iter_rows(values_only=True):
            raw_rows.append([c if c is not None else "" for c in row])

        # Find the header
        header_index = detect_header_row(raw_rows)
        if header_index is None:
            continue

        header_row = raw_rows[header_index]
        header = [str(c).strip() for c in header_row]

        # Extract data rows
        data_rows = []
        for row in raw_rows[header_index + 1:]:
            if all(not str(c).strip() for c in row):
                continue

            values = row[:len(header)]
            mapped = dict(zip(header, values))
            data_rows.append(mapped)


        tables.append({
            "sheet_name": sheet_name,
            "table_index": len(tables),
            "headers": header,
            "rows": data_rows
        })

        metadata["sheets_extracted"] += 1

    return {
        "tables": tables,
        "metadata": metadata
    }



# ---------------------------------------------------------
# week 6 PDF helpers

def page_has_text(page):
    text = page.get_text("text")
    return bool(text.strip())


def extract_text_from_page(page):
    return page.get_text("text")


def extract_images_from_page(page):
    images = []
    for img in page.get_images(full=True):
        xref = img[0]
        pix = fitz.Pixmap(page.parent, xref)

        if pix.n < 5:
            images.append(pix.pil_image())
        else:
            pix = fitz.Pixmap(fitz.csRGB, pix)
            images.append(pix.get_pil_image())

    return images


def preprocess_for_ocr(image):
    img = image.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.5)
    img = ImageEnhance.Brightness(img).enhance(1.4)
    img = img.filter(ImageFilter.SHARPEN)
    w, h = img.size
    img = img.resize((w * 3, h * 3))
    img = img.point(lambda x: 0 if x < 150 else 255, '1')
    return img


def ocr_image(image):
    processed = preprocess_for_ocr(image)
    return pytesseract.image_to_string(processed)


def parse_table_from_text(text):
    rows = []
    for line in text.split("\n"):
        cells = re.split(r"\s{2,}", line.strip())
        cells = [c for c in cells if c]
        if len(cells) >= 2:
            rows.append(cells)
    return rows


# ---------------------
# week 6 PDF extractor

def extract_pdf(path: str):
    start = time.time()
    doc = fitz.open(path)
    pages_output = []

    total_tables = 0
    total_images = 0
    ocr_used = False

    for page_num, page in enumerate(doc):
        page_result = {
            "page": page_num + 1,
            "text": None,
            "image_tables": []
        }

        page_text = extract_text_from_page(page)
        if page_text:
            page_result["text"] = page_text

        images = extract_images_from_page(page)
        total_images += len(images)

        for img in images:
            ocr_text = ocr_image(img)
            legacy_table = parse_table_from_text(ocr_text)
            if legacy_table:
                page_result["image_tables"].append({
                    "source": "pdf_ocr",
                    "page_number": page_num + 1,
                    "table_index": len(page_result["image_tables"]),
                    "headers": [],
                    "rows": legacy_table
                })

        if len(images) > 0:
            ocr_used = True

        total_tables += len(page_result["image_tables"])
        pages_output.append(page_result)

    metadata = {
        "page_count": len(doc),
        "table_count": total_tables,
        "image_count": total_images,
        "ocr_used": ocr_used,
        "file_size_kb": os.path.getsize(path) / 1024,
        "extraction_time_ms": int((time.time() - start) * 1000)
    }

    return {
        "pages": pages_output,
        "metadata": metadata
    }


# ---------------------------------
# Word docx extraction including images 

def extract_docx_text(doc):
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def extract_docx_tables(doc):
    tables_output = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)
        tables_output.append(table_data)
    return tables_output


def extract_docx_images(file_path):
    images = []
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        for file in docx_zip.namelist():
            if file.startswith("word/media/"):
                image_data = docx_zip.read(file)
                image = PILImage.open(io.BytesIO(image_data))
                images.append(image)
    return images


def extract_docx(file_path):
    start = time.time()
    doc = Document(file_path)

    text = extract_docx_text(doc)
    tables = extract_docx_tables(doc)
    images = extract_docx_images(file_path)

    metadata = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "image_count": len(images),
        "file_size_kb": os.path.getsize(file_path) / 1024,
        "extraction_time_ms": int((time.time() - start) * 1000)
    }

    all_tables = []
    text_blocks = []
    all_images = []

    for t in tables:
        all_tables.append({
            "source": "docx_table",
            "table_index": len(all_tables),
            "headers": [],
            "rows": t
        })

    if text.strip():
        text_blocks.append({
            "source": "docx_text",
            "text": text
        })

    for img in images:
        ocr_text = ocr_image(img)
        legacy_table = parse_table_from_text(ocr_text)
        if legacy_table:
            all_tables.append({
                "source": "docx_ocr",
                "table_index": len(all_tables),
                "headers": [],
                "rows": legacy_table
            })
        all_images.append(img)  # ← FIX: store actual PIL images

    return {
        "text": text_blocks,
        "tables": all_tables,
        "images": all_images,
        "metadata": metadata
    }
